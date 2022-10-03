#! python3
import numpy as np
import base64
import os.path
import docx #pip install python-docx
from os.path import join

# convert data to binary
def dataToBin(data):
    if type(data) == str:
        return ''.join([format(ord(i), "08b") for i in data])
    elif type(data) == bytes or type(data) == np.ndarray:
        return [format(i, "08b") for i in data]
    elif type(data) == int or type(data) == np.uint8:
        return format(data, "08b")
    else:
        # Input type not able to be changed to binary
        raise TypeError("Input type not supported")

# Function will open docx file and put into string
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    str ='\n'.join(fullText)
    return ''.join([i if ord(i) < 128 else ' ' for i in str])


# Function will create a docx and save the data into the file
def createDoc(text):
    path=os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop'), # set path to open desktop
    name = 'encoded_file.docx'
    path = ''.join(path)
    output_path = join(path, name)
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(output_path)
    return(output_path)

def encode_docx(cover, payload, lsb):
    coverString = getText(cover)
    payloadPath = os.path.join(os.getcwd(), payload)

    # Opens the payload file and encodes it into a Base64 format
    payloadB64 = ''
    with open(payloadPath, "rb") as payload:
        payloadB64 = base64.b64encode(payload.read())

    # Find extension of payload file
    payloadExt = os.path.splitext(payloadPath)[1]

    # Add padding to file extention untill there is 10 characters
    while len(payloadExt) < 10:
        payloadExt += '@'

    # Converts the encoded Base64 payload into a string and adds the file extention and the delimeter to indicate the end of the file
    payloadEncode = str(payloadB64) + payloadExt + "%$&@"

    # Converts the encoded payload string to binary format
    payloadBin = dataToBin(payloadEncode)

    coverString = coverString.replace('\n','')
    coverString = coverString.replace('\t','')
    coverString = coverString.replace('\r','')
    # Convert cover string to binary format
    coverBin = dataToBin(coverString)

    # splitting coverBin to bytes (8 bits per 1 byte) > xxxxxxxx xxxxxxxx xxxxxxxx
    coverBinByte = []
    n  = 8
    for index in range(0, len(coverBin), n):
        coverBinByte.append(coverBin[index : index + n])
    
    # check if coverData is big enough for payloadData
    if len(payloadBin) / lsb > len(coverBin) / 8:
        raise ValueError("[!] Insufficient bytes, need bigger cover document")
    print ("[*] Encoding data... \n")

    # Get the length of the encoded payload binary
    payloadLen = len(payloadBin)
    # Setting a local variable to point to the binary value to encode in each iteration
    dataIndex = 0

    for i,byte in enumerate(coverBinByte):
        # If there is still more data to store
        if dataIndex < payloadLen:
            tempBin = ''
            # Remove the last few digits of the byte based on the number of LSB used
            byte = byte[:-lsb]
            # Loop and store the data to hide into a temp variable
            for x in range(lsb):
                tempBin += payloadBin[dataIndex]
                dataIndex += 1
                # Break the loop if all the data have been hidden
                if dataIndex == payloadLen:
                    break
            # Pad 0 to the front of the temp binary if the length of the temp binary is less than the number of LSB used
            if len(tempBin) < lsb:
                tempBin = tempBin.ljust(lsb, '0')
            # Adds the temp binary to the byte and change it into an integer and assign it to the back to the array of bytes
            coverBinByte[i] = byte + tempBin

    # join all the bit together
    newEncodedBin = "".join(coverBinByte)
    #print("newEncodedBin: ", newEncodedBin)

    # splitting by 8-bits  
    allBytes = [newEncodedBin[i: i + 8] for i in range(0, len(newEncodedBin), 8)]  
    # converting from bits to characters  
    encodedText = ""  
    for bytes in allBytes:  
        encodedText += chr(int(bytes, 2))  
    
    #print("\nEncoded Text(Stego):", encodedText)
    #print("Length:", len(encodedText), "\n")
    print("########### Encoding Successful ###########\n")
    #print(type(encodedText))
    #print(encodedText)
    output_path = createDoc(encodedText)
    return(output_path)


# Function to decode the stego and find the payload
def decode_docx(stegoFile, lsb):
    print("[*] Decoding data... \n")
    
    # Get path of stego file
    stegoPath = os.path.join(os.getcwd(), stegoFile)

    stegoString = getText(stegoPath)

    # splitting stegoBytes to binary (8 bits per 1 byte) > xxxxxxxx xxxxxxxx xxxxxxxx
    stegoByteList = []
    for i in range(0, len(stegoString)):
        stegoByteList.append(dataToBin(stegoString[i]))


    # Extract the LSB based on the value defined by the user
    binaryData = ''
    for i,byte in enumerate(stegoByteList):
        binaryData += byte[-lsb:]

    # splitting by 8-bits  (1 byte) [xxxxxxxx,xxxxxxxx, xxxxxxxx]
    allBytes = [binaryData[i: i + 8] for i in range(0, len(binaryData), 8)]

    # converting from bits to characters  
    decodedData = ""  
    for bytes in allBytes:  
        decodedData += chr(int(bytes, 2))  
        # checking if we have reached the delimiter which is "%$&@"  
        if decodedData[-4:] == "%$&@":  
            break
        
    # Removes delimeter
    decodedData = decodedData[:-4]
    # Get the file extension and removes it
    fileExt = decodedData[-10:]
    decodedData = decodedData[:-10]
    # Remove padding from file extention
    fileExt = fileExt.replace("@", "")
    
    # Converts the hidden data back to a file from a Base64 format
    decodedData = base64.b64decode(eval(decodedData))

    # Write the decoded data back to a file and saves it
    path=os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop'), # set path to open desktop
    name = 'decoded_output' + fileExt
    path = ''.join(path)
    output_path = join(path, name)
    with open(output_path, "wb") as outFile:
        outFile.write(decodedData)
    
    return(output_path)


'''##################################################################################################'''

# cover = "appleTNC.docx"
# payload = "payload.txt"

# encode(cover, payload, 1)
# decode("plesework.docx", 1)

